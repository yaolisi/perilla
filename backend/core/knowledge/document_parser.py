"""
文档解析器
支持 PDF, DOCX, TXT/MD 格式
"""
from __future__ import annotations

from typing import List, Dict, Any, Optional
from pathlib import Path
from log import logger


class ParsedPage:
    """解析后的页面"""
    def __init__(self, page: int, text: str) -> None:
        self.page = page
        self.text = text.strip()


class ParsedDocument:
    """解析后的文档"""
    def __init__(
        self,
        doc_id: str,
        filename: str,
        pages: List[ParsedPage],
        doc_type: str,
    ) -> None:
        self.doc_id = doc_id
        self.filename = filename
        self.pages = pages
        self.doc_type = doc_type
    
    def get_full_text(self) -> str:
        """获取完整文本"""
        return "\n\n".join(page.text for page in self.pages)
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            "doc_id": self.doc_id,
            "filename": self.filename,
            "pages": [{"page": p.page, "text": p.text} for p in self.pages],
            "doc_type": self.doc_type,
        }


class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    def parse(file_path: Path, doc_type: Optional[str] = None) -> ParsedDocument:
        """
        解析文档
        
        Args:
            file_path: 文件路径
            doc_type: 文档类型（pdf/docx/txt/md），如果不提供则从扩展名推断
            
        Returns:
            ParsedDocument
            
        Raises:
            ValueError: 不支持的文件类型
            Exception: 解析失败
        """
        if doc_type is None:
            ext = file_path.suffix.lower()
            doc_type = ext[1:] if ext else "txt"
        
        doc_type = doc_type.lower()
        
        if doc_type == "pdf":
            return DocumentParser._parse_pdf(file_path)
        elif doc_type == "docx":
            return DocumentParser._parse_docx(file_path)
        elif doc_type in ["txt", "md"]:
            return DocumentParser._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")
    
    @staticmethod
    def _parse_pdf(file_path: Path) -> ParsedDocument:
        """解析 PDF"""
        try:
            import pdfplumber  # type: ignore[import-untyped]
        except ImportError:
            try:
                import pypdf
            except ImportError:
                raise ImportError("Please install pdfplumber or pypdf: pip install pdfplumber")
        
        pages = []
        doc_id = file_path.stem
        
        try:
            import pdfplumber  # type: ignore[import-untyped]
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text and text.strip():
                        pages.append(ParsedPage(page=page_num, text=text))
        except ImportError:
            # Fallback to pypdf
            import pypdf  # type: ignore[import-untyped]
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    text = page.extract_text()
                    if text and text.strip():
                        pages.append(ParsedPage(page=page_num, text=text))
        
        if not pages:
            raise ValueError("PDF contains no extractable text (may be scanned image)")
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=file_path.name,
            pages=pages,
            doc_type="pdf",
        )
    
    @staticmethod
    def _parse_docx(file_path: Path) -> ParsedDocument:
        """解析 DOCX"""
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError:
            raise ImportError("Please install python-docx: pip install python-docx")
        
        doc_id = file_path.stem
        pages = []
        
        doc = Document(file_path)
        
        # DOCX 没有明确的页概念，将所有内容作为 page 0
        full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        
        if full_text.strip():
            pages.append(ParsedPage(page=0, text=full_text))
        
        if not pages:
            raise ValueError("DOCX contains no text")
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=file_path.name,
            pages=pages,
            doc_type="docx",
        )
    
    @staticmethod
    def _parse_text(file_path: Path) -> ParsedDocument:
        """解析 TXT/MD"""
        doc_id = file_path.stem
        doc_type = file_path.suffix.lower()[1:] if file_path.suffix else "txt"
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(file_path, 'r', encoding='gbk') as f:
                text = f.read()
        
        if not text.strip():
            raise ValueError("File is empty")
        
        pages = [ParsedPage(page=0, text=text)]
        
        return ParsedDocument(
            doc_id=doc_id,
            filename=file_path.name,
            pages=pages,
            doc_type=doc_type,
        )
